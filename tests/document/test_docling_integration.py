"""End-to-end inventory against the real Docling backend.

Skipped unless the optional `document` extra is installed, so the default GPU-free
test run stays free of a heavy dependency. The fixture is generated rather than
committed: a binary fixture in the repository would be one more artifact to keep
in step with the backend.
"""

from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docling", reason="requires the optional document extra")
pytest.importorskip("docx", reason="requires the optional document extra")

from docx import Document as DocxDocument  # noqa: E402

from document.docling_adapter import DoclingDocumentParser  # noqa: E402
from document.inventory import build_inventory  # noqa: E402

pytestmark = pytest.mark.docling


@pytest.fixture(scope="module")
def report(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """A minimal report carrying one borehole log heading and one table."""
    document = DocxDocument()
    document.add_heading("Ground Investigation Report", level=1)
    document.add_paragraph("Borehole log BH-07 recorded at chainage 1200.")

    table = document.add_table(rows=2, cols=3)
    for column, heading in enumerate(("From", "To", "Lithology")):
        table.cell(0, column).text = heading
    for column, value in enumerate(("0.0", "3.2", "Fill")):
        table.cell(1, column).text = value

    path = tmp_path_factory.mktemp("docling") / "report.docx"
    document.save(path)
    return path


def test_real_backend_produces_a_cited_inventory(report: Path) -> None:
    parsed = DoclingDocumentParser().parse(report)
    inventory = build_inventory("doc-real", "sha-real", parsed)

    assert inventory.source_format == "docx"
    assert inventory.page_count >= 1
    assert any(page.has_text for page in inventory.pages)

    text = " ".join(page.text for page in inventory.pages).casefold()
    assert "bh-07" in text

    # Every routed region must cite a real page of this document.
    page_numbers = {page.page_number for page in inventory.pages}
    for figure in inventory.candidates():
        evidence = inventory.evidence_for(figure)
        assert evidence.document_id == "doc-real"
        assert evidence.page_number in page_numbers


def test_real_backend_routes_the_borehole_log_text(report: Path) -> None:
    parsed = DoclingDocumentParser().parse(report)
    inventory = build_inventory("doc-real", "sha-real", parsed)

    # The heading text names a borehole log, so any region on that page is routed
    # there rather than being left as a generic table.
    assert inventory.candidates("borehole_log") or not inventory.candidates()


def _build_pdf(pages: list[list[tuple[int, int, int, str]]]) -> bytes:
    """A minimal multi-page PDF with a real text layer, built without a dependency."""
    objects: list[bytes] = []

    def add(body: bytes) -> int:
        objects.append(body)
        return len(objects)

    content_ids = []
    for lines in pages:
        parts = [b"BT"]
        for x, y, size, text in lines:
            escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
            parts.append(f"/F1 {size} Tf {x} {y} Td ({escaped}) Tj".encode("latin-1"))
        parts.append(b"ET")
        stream = b"\n".join(parts)
        content_ids.append(
            add(b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream))
        )

    font_id = add(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    pages_id = len(objects) + len(pages) + 1
    page_ids = [
        add(
            b"<< /Type /Page /Parent %d 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 %d 0 R >> >> /Contents %d 0 R >>"
            % (pages_id, font_id, content_id)
        )
        for content_id in content_ids
    ]
    kids = b" ".join(b"%d 0 R" % pid for pid in page_ids)
    add(b"<< /Type /Pages /Kids [%s] /Count %d >>" % (kids, len(page_ids)))
    catalog_id = add(b"<< /Type /Catalog /Pages %d 0 R >>" % pages_id)

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for index, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n%s\nendobj\n" % (index, body)
    xref_at = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objects) + 1)
    for offset in offsets:
        out += b"%010d 00000 n \n" % offset
    out += b"trailer\n<< /Size %d /Root %d 0 R >>\nstartxref\n%d\n%%%%EOF\n" % (
        len(objects) + 1,
        catalog_id,
        xref_at,
    )
    return bytes(out)


@pytest.fixture(scope="module")
def pdf_report(tmp_path_factory: pytest.TempPathFactory) -> Path:
    path = tmp_path_factory.mktemp("docling-pdf") / "report.pdf"
    path.write_bytes(
        _build_pdf(
            [
                [(72, 720, 18, "Ground Investigation Report")],
                [
                    (72, 720, 16, "Borehole Log BH-07"),
                    (72, 690, 11, "0.0 - 3.2 m Fill"),
                ],
            ]
        )
    )
    return path


def test_pdf_backend_reports_real_pagination(pdf_report: Path) -> None:
    """Assert the structure a PDF must yield, not which text the model recovers.

    Docling attaches text to regions its layout model detects. On this synthetic
    fixture that detection is not stable: repeated runs of the identical file and
    options recovered page 1 only, page 2 only, or both. Two floating lines with no
    document structure are out of distribution for a model trained on real reports,
    so asserting specific extracted text here would test the model's stability on
    degenerate input rather than this adapter.

    The page structure, by contrast, was identical on every run, and it is what the
    adapter is responsible for. Text completeness against real-world PDFs is a
    separate question, recorded as an open risk in PLAN.md.
    """
    parsed = DoclingDocumentParser().parse(pdf_report)
    inventory = build_inventory("doc-pdf", "sha-pdf", parsed)

    # Unlike DOCX, a PDF has fixed pages, so the page numbers are the source's.
    assert inventory.has_source_pagination is True
    assert inventory.source_format == "pdf"
    assert inventory.page_count == 2
    assert [page.page_number for page in inventory.pages] == [1, 2]


def test_pdf_content_is_attributed_only_to_real_pages(pdf_report: Path) -> None:
    """Whatever the backend recovers must be cited to a page that exists."""
    parsed = DoclingDocumentParser().parse(pdf_report)
    inventory = build_inventory("doc-pdf", "sha-pdf", parsed)

    page_numbers = {page.page_number for page in inventory.pages}
    for figure in inventory.candidates():
        assert figure.page_number in page_numbers
        assert inventory.evidence_for(figure).page_number in page_numbers
        if figure.bbox is not None:
            left, top, right, bottom = figure.bbox
            assert right >= left and bottom >= top
